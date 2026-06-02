"""Input-validation and routing tests for WordRuntimeEditTool.

A fake backend is injected via the global session manager so the tool path is
exercised without requiring Word/COM.
"""

from __future__ import annotations

import asyncio
import json
from pathlib import Path

import docx
import pytest

from agent.core.loop import LoopConfig, LoopContext
from agent.core.word_runtime import (
    HeadingInfo,
    WordRuntimeRequest,
    WordRuntimeResult,
    WordStructure,
    get_session_manager,
)
from agent.tools_v2.word_runtime_tool import WordRuntimeEditTool


class StructuralFakeBackend:
    name = "fake-structural"

    def __init__(self) -> None:
        self.connected = False
        self.requests: list[WordRuntimeRequest] = []

    def connect(self) -> None:
        self.connected = True

    def shutdown(self) -> None:
        self.connected = False

    def is_alive(self) -> bool:
        return self.connected

    def get_structure(self, path: Path) -> WordStructure:
        return WordStructure()

    def apply(self, request: WordRuntimeRequest) -> WordRuntimeResult:
        self.requests.append(request)
        return WordRuntimeResult(
            path=str(request.path),
            backup_path=None,
            ops_applied=len(request.ops),
            structure_before={
                "headings": [{"text": "Chapter 1", "level": 1, "paragraph_index": 0}],
                "has_toc_field": False,
            },
            structure_after={
                "headings": [
                    {"text": "Chapter 1", "level": 1, "paragraph_index": 0},
                    {"text": "Chapter 2", "level": 1, "paragraph_index": 5},
                ],
                "has_toc_field": True,
            },
        )


@pytest.fixture
def fake_backend(tmp_path):
    backend = StructuralFakeBackend()
    manager = get_session_manager(lambda: backend)
    manager.shutdown_all()
    manager.configure_factory(lambda: backend)
    yield backend
    manager.shutdown_all()


def _ctx() -> LoopContext:
    return LoopContext(config=LoopConfig())


def _docx(path: Path, headings=("Chapter 1",)) -> None:
    document = docx.Document()
    for h in headings:
        document.add_heading(h, level=1)
        document.add_paragraph("body text")
    document.save(path)


def test_tool_rejects_paragraph_index_anchor(tmp_path, fake_backend):
    target = tmp_path / "doc.docx"
    _docx(target)
    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{
            "op": "set_heading_text",
            "paragraph_index": 0,
            "new_text": "X",
        }],
    }, _ctx()))
    assert res.is_error is True
    assert "paragraph_index" in res.content


def test_tool_rejects_unsupported_op(tmp_path, fake_backend):
    target = tmp_path / "doc.docx"
    _docx(target)
    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{"op": "replace_text", "old": "a", "new": "b"}],
    }, _ctx()))
    assert res.is_error is True
    assert "structural" in res.content.lower()


def test_tool_routes_refresh_fields(tmp_path, fake_backend):
    target = tmp_path / "doc.docx"
    _docx(target)
    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{"op": "refresh_fields"}],
    }, _ctx()))
    assert res.is_error is False
    payload = json.loads(res.content)
    assert payload["ops_applied"] == 1
    assert payload["structure_after"]["has_toc_field"] is True
    assert fake_backend.requests
    assert fake_backend.requests[0].ops[0].op.value == "refresh_fields"


def test_tool_anchor_required_for_set_heading(tmp_path, fake_backend):
    target = tmp_path / "doc.docx"
    _docx(target)
    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{"op": "set_heading_text", "new_text": "X"}],
    }, _ctx()))
    assert res.is_error is True
    assert "anchor" in res.content.lower()


def test_tool_factory_wires_word_runtime_edit():
    from agent.tools_v2.factory import build_tool

    tool = build_tool("WordRuntimeEdit")
    assert tool.name == "WordRuntimeEdit"


def test_tool_rejects_non_docx(tmp_path, fake_backend):
    target = tmp_path / "doc.txt"
    target.write_text("hi", encoding="utf-8")
    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{"op": "refresh_fields"}],
    }, _ctx()))
    assert res.is_error is True
    assert ".docx" in res.content


def test_tool_registers_artifact_manifest_after_success(tmp_path, fake_backend):
    """After WordRuntimeEdit succeeds, the artifact registry must hold the
    post-mutation manifest for the touched path under this conversation."""
    from agent.core.artifact_context import (
        get_registry,
        reset_registry,
    )

    reset_registry("conv-manifest")
    target = tmp_path / "doc.docx"
    _docx(target)

    ctx = LoopContext(config=LoopConfig())
    ctx.scratch["conversation_id"] = "conv-manifest"

    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{"op": "refresh_fields"}],
    }, ctx))
    assert res.is_error is False

    record = get_registry("conv-manifest").get(str(target))
    assert record is not None, "artifact registry was not updated after edit"
    # structure_after.has_toc_field == True → manifest.has_toc reflects post-edit state.
    assert record.manifest.has_toc is True
    # Headings reflect structure_after (2 chapters), not structure_before (1 chapter).
    assert len(record.manifest.headings) == 2


def test_tool_registers_artifact_manifest_after_success(tmp_path, fake_backend):
    """After WordRuntimeEdit succeeds, the artifact registry must hold the
    post-mutation manifest for the touched path under this conversation."""
    from agent.core.artifact_context import (
        get_registry,
        reset_registry,
    )

    reset_registry("conv-manifest")
    target = tmp_path / "doc.docx"
    _docx(target)

    ctx = LoopContext(config=LoopConfig())
    ctx.scratch["conversation_id"] = "conv-manifest"

    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{"op": "refresh_fields"}],
    }, ctx))
    assert res.is_error is False

    record = get_registry("conv-manifest").get(str(target))
    assert record is not None, "artifact registry was not updated after edit"
    # structure_after.has_toc_field == True → manifest.has_toc reflects post-edit state.
    assert record.manifest.has_toc is True
    # Headings reflect structure_after (2 chapters), not structure_before (1 chapter).
    assert len(record.manifest.headings) == 2


def test_tool_registers_artifact_manifest_after_success(tmp_path, fake_backend):
    """After WordRuntimeEdit succeeds, the artifact registry must hold the
    post-mutation manifest for the touched path under this conversation."""
    from agent.core.artifact_context import (
        get_registry,
        reset_registry,
    )

    reset_registry("conv-manifest")
    target = tmp_path / "doc.docx"
    _docx(target)

    ctx = LoopContext(config=LoopConfig())
    ctx.scratch["conversation_id"] = "conv-manifest"

    tool = WordRuntimeEditTool()
    res = asyncio.run(tool.run({
        "path": str(target),
        "ops": [{"op": "refresh_fields"}],
    }, ctx))
    assert res.is_error is False

    record = get_registry("conv-manifest").get(str(target))
    assert record is not None, "artifact registry was not updated after edit"
    # structure_after.has_toc_field == True → manifest.has_toc reflects post-edit state.
    assert record.manifest.has_toc is True
    # Headings reflect structure_after (2 chapters), not structure_before (1 chapter).
    assert len(record.manifest.headings) == 2
