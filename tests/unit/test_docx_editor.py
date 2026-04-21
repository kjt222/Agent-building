import tempfile
import unittest
from pathlib import Path

import docx

from agent.tools import apply_docx_ops


def _make_split_run_paragraph(document, fragments_and_bold):
    """Append a paragraph composed of multiple runs with individual formatting."""
    paragraph = document.add_paragraph()
    for text, bold in fragments_and_bold:
        run = paragraph.add_run(text)
        run.bold = bold
    return paragraph


class TestDocxEditor(unittest.TestCase):
    def test_replace_and_append(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            document = docx.Document()
            document.add_paragraph("hello world")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "table world"
            document.save(path)

            ops = [
                {"op": "replace_text", "old": "world", "new": "there"},
                {"op": "append_paragraph", "text": "append"},
            ]
            result = apply_docx_ops(path, ops)
            self.assertGreaterEqual(result.replacements, 2)

            updated = docx.Document(path)
            body_text = "\n".join(p.text for p in updated.paragraphs if p.text)
            self.assertIn("hello there", body_text)
            self.assertIn("append", body_text)
            self.assertEqual(updated.tables[0].cell(0, 0).text, "table there")

    def test_single_run_edit_preserves_format(self) -> None:
        """Baseline: target text inside one run — its rPr must be preserved."""
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            document = docx.Document()
            _make_split_run_paragraph(
                document,
                [("Before ", False), ("bold-target", True), (" after", False)],
            )
            document.save(path)

            result = apply_docx_ops(
                path, [{"op": "replace_text", "old": "bold-target", "new": "BOLD"}]
            )
            self.assertEqual(result.replacements, 1)
            self.assertEqual(result.cross_run_merges, 0)

            updated = docx.Document(path)
            para = updated.paragraphs[0]
            self.assertEqual(para.text, "Before BOLD after")
            # The run that originally held "bold-target" still exists and is bold.
            bold_runs = [r for r in para.runs if r.bold]
            self.assertTrue(any("BOLD" in r.text for r in bold_runs))

    def test_cross_run_merge_preserves_surrounding_format(self) -> None:
        """Cross-run match: first spanned run keeps its rPr; runs outside the
        span retain their own formatting."""
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            document = docx.Document()
            # "Hello " (plain) | "Wor" (bold) | "ld!" (italic)
            # Target "World" spans run 2 and run 3.
            paragraph = document.add_paragraph()
            r1 = paragraph.add_run("Hello ")
            r2 = paragraph.add_run("Wor")
            r2.bold = True
            r3 = paragraph.add_run("ld!")
            r3.italic = True
            document.save(path)

            result = apply_docx_ops(
                path, [{"op": "replace_text", "old": "World", "new": "Universe"}]
            )
            self.assertEqual(result.replacements, 1)
            self.assertEqual(result.cross_run_merges, 1)

            updated = docx.Document(path)
            para = updated.paragraphs[0]
            self.assertEqual(para.text, "Hello Universe!")
            # First run still plain (not bold/italic)
            self.assertFalse(para.runs[0].bold)
            self.assertFalse(para.runs[0].italic)
            # The run that originally carried "Wor" absorbed "Universe" and is still bold.
            bold_runs = [r for r in para.runs if r.bold]
            self.assertTrue(any("Universe" in r.text for r in bold_runs))

    def test_fail_on_cross_run_raises(self) -> None:
        """Strict mode: cross-run matches should raise rather than rewrite."""
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Hello ")
            paragraph.add_run("Wor").bold = True
            paragraph.add_run("ld!").italic = True
            document.save(path)

            with self.assertRaises(ValueError):
                apply_docx_ops(
                    path,
                    [{"op": "replace_text", "old": "World", "new": "Universe"}],
                    fail_on_cross_run=True,
                )
