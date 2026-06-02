"""Tests for the Office L2 oracle (P14.2.3)."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from agent.acceptance.oracles.office import OfficeOracle
from agent.acceptance.oracle import get_oracle


def test_oracle_is_registered():
    from agent.acceptance import oracles  # noqa: F401
    assert get_oracle("office") is not None


def test_empty_input_returns_unknown():
    rep = OfficeOracle().check([])
    assert rep.verdict == "unknown"


def test_unsupported_extension_is_unknown():
    rep = OfficeOracle().check([Path("nope.txt")])
    # No probable office file → unknown (not fail — we don't claim
    # ownership over txts)
    assert rep.verdict == "unknown"


def test_missing_docx_fails(tmp_path):
    p = tmp_path / "ghost.docx"
    rep = OfficeOracle().check([p])
    assert rep.verdict == "fail"
    assert any("missing" in f for f in rep.findings)


def test_real_docx_with_content_passes(tmp_path):
    try:
        from docx import Document  # type: ignore
    except Exception:
        pytest.skip("python-docx not installed")
    p = tmp_path / "real.docx"
    doc = Document()
    doc.add_paragraph("Hello P14.2.3.")
    doc.save(str(p))
    rep = OfficeOracle().check([p])
    assert rep.verdict == "pass", rep.findings
    files_ev = rep.evidence["files"][str(p)]
    assert files_ev["non_empty_paragraph_count"] >= 1


def test_empty_docx_fails(tmp_path):
    try:
        from docx import Document  # type: ignore
    except Exception:
        pytest.skip("python-docx not installed")
    p = tmp_path / "empty.docx"
    doc = Document()
    # Don't add anything — default empty doc has 0 non-empty paragraphs.
    doc.save(str(p))
    rep = OfficeOracle().check([p])
    assert rep.verdict == "fail"
    assert any("no non-empty paragraphs" in f for f in rep.findings)


def test_xlsx_with_content_passes(tmp_path):
    try:
        from openpyxl import Workbook  # type: ignore
    except Exception:
        pytest.skip("openpyxl not installed")
    p = tmp_path / "real.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "hello"
    ws["B1"] = 42
    wb.save(str(p))
    rep = OfficeOracle().check([p])
    assert rep.verdict == "pass"
    assert rep.evidence["files"][str(p)]["non_empty_cell_count"] >= 2
