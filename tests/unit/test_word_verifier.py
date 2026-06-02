"""Toy verifier tests with hand-built .docx fixtures.

Fixtures use python-docx to build documents with specific structural
properties so the verifier rules can be exercised without a Word runtime.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import docx

from agent.tools_v2.word_verifier import verify_word_document


def _good_doc(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("First chapter body.")
    document.add_heading("Chapter 2", level=1)
    document.add_paragraph("Second chapter body.")
    document.save(path)


def _doc_with_orphan_in_table(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Chapter 1", level=1)
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    rogue = cell.paragraphs[0]
    rogue.text = "Rogue Heading"
    rogue.style = "Heading 1"
    document.add_paragraph("normal body")
    document.save(path)


def _doc_with_empty_heading(path: Path) -> None:
    document = docx.Document()
    document.add_heading("", level=1)
    document.add_paragraph("body")
    document.save(path)


def _inject_unrefreshed_toc(path: Path) -> None:
    """Add a TOC field block with empty cache so verifier can flag it."""
    document = docx.Document()
    document.add_heading("Chapter 1", level=1)
    para = document.add_paragraph()
    run = para.add_run()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h')
    inner_run = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = ""
    inner_run.append(inner_t)
    fld.append(inner_run)
    run._element.addnext(fld)

    document.add_heading("Chapter 2", level=1)
    document.save(path)


def test_verifier_passes_clean_doc(tmp_path):
    target = tmp_path / "good.docx"
    _good_doc(target)
    result = verify_word_document(target)
    assert result.passed is True, result.violations
    assert result.metrics["heading_count"] == 2


def test_verifier_flags_heading_inside_table(tmp_path):
    target = tmp_path / "rogue.docx"
    _doc_with_orphan_in_table(target)
    result = verify_word_document(target)
    assert result.passed is False
    assert "heading_inside_table_cell" in result.violations
    assert any("table cell" in hint for hint in result.repair_hints)


def test_verifier_flags_empty_heading(tmp_path):
    target = tmp_path / "empty.docx"
    _doc_with_empty_heading(target)
    result = verify_word_document(target)
    assert result.passed is False
    assert "empty_heading_paragraphs" in result.violations


def test_verifier_flags_unrefreshed_toc(tmp_path):
    target = tmp_path / "stale.docx"
    _inject_unrefreshed_toc(target)
    result = verify_word_document(target)
    assert result.passed is False
    assert "toc_cache_empty" in result.violations or "all_fields_unrefreshed" in result.violations
    assert any("refresh" in hint for hint in result.repair_hints)


def test_verifier_missing_file(tmp_path):
    result = verify_word_document(tmp_path / "nope.docx")
    assert result.passed is False
    assert any(v.startswith("missing_file") for v in result.violations)


def test_verifier_require_toc_when_absent(tmp_path):
    target = tmp_path / "no-toc.docx"
    _good_doc(target)
    result = verify_word_document(target, require_toc=True)
    assert result.passed is False
    assert "missing_toc_field" in result.violations
