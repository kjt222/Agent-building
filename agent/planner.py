from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, Iterable, Optional

import yaml

from .models import ModelAdapter


def extract_docx_preview(
    path: Path, max_paragraphs: int = 200, max_chars: int = 12000
) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx not installed. Install `python-docx`.") from exc
    document = docx.Document(str(path))
    lines = []
    used = 0
    para_index = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        para_index += 1
        style_name = paragraph.style.name if paragraph.style else ""
        prefix = f"[{para_index}]"
        if style_name:
            line = f"{prefix} ({style_name}) {text}"
        else:
            line = f"{prefix} {text}"
        if used + len(line) + 1 > max_chars:
            break
        lines.append(line)
        used += len(line) + 1
        if para_index >= max_paragraphs:
            break
    if used < max_chars:
        for t_index, table in enumerate(document.tables, start=1):
            for r_index, row in enumerate(table.rows, start=1):
                for c_index, cell in enumerate(row.cells, start=1):
                    cell_text = " ".join(p.text for p in cell.paragraphs if p.text).strip()
                    if not cell_text:
                        continue
                    line = f"[T{t_index}R{r_index}C{c_index}] {cell_text}"
                    if used + len(line) + 1 > max_chars:
                        return "\n".join(lines)
                    lines.append(line)
                    used += len(line) + 1
    return "\n".join(lines)


def extract_xlsx_preview(
    path: Path,
    max_rows: int = 20,
    max_cols: int = 10,
    max_sheets: int = 3,
    max_chars: int = 12000,
    cell_chars: int = 40,
) -> str:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("openpyxl not installed. Install `openpyxl`.") from exc
    workbook = openpyxl.load_workbook(path, data_only=False, read_only=True)
    lines = []
    used = 0
    for s_index, sheet in enumerate(workbook.worksheets, start=1):
        if s_index > max_sheets:
            break
        header = f"Sheet: {sheet.title} range: {sheet.calculate_dimension()}"
        if used + len(header) + 1 > max_chars:
            break
        lines.append(header)
        used += len(header) + 1
        for row_idx in range(1, max_rows + 1):
            values = []
            for col_idx in range(1, max_cols + 1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                value = cell.value
                if value is None:
                    display = ""
                else:
                    display = str(value)
                if len(display) > cell_chars:
                    display = f"{display[:cell_chars]}..."
                values.append(f"{cell.coordinate}={display}")
            line = f"Row {row_idx}: " + " | ".join(values).rstrip()
            if used + len(line) + 1 > max_chars:
                break
            lines.append(line)
            used += len(line) + 1
        lines.append("")
        used += 1
        if used >= max_chars:
            break
    return "\n".join(lines).strip()


def _extract_code_block(text: str) -> str:
    match = re.search(r"```(?:json|yaml)?\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return text.strip()


def parse_plan_text(text: str) -> Dict[str, Any]:
    payload = _extract_code_block(text)
    try:
        data = yaml.safe_load(payload)
    except yaml.YAMLError as exc:
        raise ValueError("Failed to parse plan output") from exc
    if not isinstance(data, dict):
        raise ValueError("Plan output must be a mapping")
    return data


def _validate_ops(ops: Iterable[dict], allowed: set[str]) -> None:
    for op in ops:
        op_type = op.get("op")
        if op_type not in allowed:
            raise ValueError(f"Unsupported op: {op_type}")


def validate_docx_plan(plan: Dict[str, Any]) -> None:
    ops = plan.get("ops")
    if not isinstance(ops, list) or not ops:
        raise ValueError("Docx plan must include ops list")
    _validate_ops(ops, {"replace_text", "append_paragraph", "add_heading"})
    for op in ops:
        op_type = op.get("op")
        if op_type == "replace_text" and ("old" not in op or "new" not in op):
            raise ValueError("replace_text requires old/new")
        if op_type in {"append_paragraph", "add_heading"} and "text" not in op:
            raise ValueError(f"{op_type} requires text")


def validate_xlsx_plan(plan: Dict[str, Any]) -> None:
    ops = plan.get("ops")
    if not isinstance(ops, list) or not ops:
        raise ValueError("Xlsx plan must include ops list")
    _validate_ops(
        ops,
        {
            "set_cell",
            "fill_formula",
            "insert_columns",
            "delete_columns",
            "insert_rows",
            "delete_rows",
            "set_auto_filter",
            "sort_range",
            "pivot",
        },
    )
    for op in ops:
        op_type = op.get("op")
        if op_type == "set_cell" and "cell" not in op:
            raise ValueError("set_cell requires cell")
        if op_type == "fill_formula" and ("range" not in op or "formula" not in op):
            raise ValueError("fill_formula requires range/formula")


def build_docx_prompt(instruction: str, preview: str, file_name: str) -> str:
    return (
        "You are an assistant that prepares DOCX edit plans.\n"
        "Return ONLY YAML or JSON with this schema:\n"
        "file: <string>\n"
        "ops: list of operations\n"
        "Supported ops:\n"
        "- op: replace_text, old: <exact text>, new: <replacement>, count: <optional int>\n"
        "- op: append_paragraph, text: <string>, style: <optional style>\n"
        "- op: add_heading, text: <string>, level: <1-6>\n"
        "Rules:\n"
        "- Use exact text from the preview for replace_text.old.\n"
        "- Keep edits minimal and avoid reformatting.\n"
        "- Only use the file shown.\n"
        f"Target file: {file_name}\n"
        f"Preview:\n{preview}\n\n"
        f"Instruction: {instruction}\n"
    )


def build_xlsx_prompt(instruction: str, preview: str, file_name: str) -> str:
    return (
        "You are an assistant that prepares XLSX edit plans.\n"
        "Return ONLY YAML or JSON with this schema:\n"
        "file: <string>\n"
        "sheet: <optional default sheet>\n"
        "ops: list of operations\n"
        "Supported ops:\n"
        "- op: set_cell, cell: <A1>, value: <value> OR formula: <formula>\n"
        "- op: fill_formula, range: <A1:C10>, formula: <template with {row} {col}>\n"
        "- op: insert_rows/delete_rows, index: <row>, amount: <int>\n"
        "- op: insert_columns/delete_columns, index: <col>, amount: <int>\n"
        "- op: set_auto_filter, range: <A1:C10>\n"
        "- op: sort_range, range: <A1:C10>, key: <A or column index>, header: <true/false>, ascending: <true/false>\n"
        "Rules:\n"
        "- Only reference sheet names that exist in the preview.\n"
        "- Use explicit cell references and ranges.\n"
        "- Keep changes minimal.\n"
        f"Target file: {file_name}\n"
        f"Preview:\n{preview}\n\n"
        f"Instruction: {instruction}\n"
    )


def generate_docx_plan(
    llm: ModelAdapter,
    instruction: str,
    preview: str,
    file_path: Path,
) -> Dict[str, Any]:
    prompt = build_docx_prompt(instruction, preview, file_path.name)
    raw = llm.chat(prompt)
    plan = parse_plan_text(raw)
    plan["file"] = str(file_path)
    validate_docx_plan(plan)
    return plan


def generate_xlsx_plan(
    llm: ModelAdapter,
    instruction: str,
    preview: str,
    file_path: Path,
) -> Dict[str, Any]:
    prompt = build_xlsx_prompt(instruction, preview, file_path.name)
    raw = llm.chat(prompt)
    plan = parse_plan_text(raw)
    plan["file"] = str(file_path)
    validate_xlsx_plan(plan)
    return plan


def write_plan(plan: Dict[str, Any], path: Path) -> None:
    payload = yaml.safe_dump(plan, sort_keys=False, allow_unicode=False)
    path.write_text(payload, encoding="utf-8")
