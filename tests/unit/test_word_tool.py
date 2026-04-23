from __future__ import annotations

import asyncio
import json
import zipfile

import docx
from docx.shared import Pt

from agent.core.loop import LoopConfig, LoopContext, PermissionLevel
from agent.tools_v2.word_tool import WordEditTool, WordReadTool, word_toolset


def _ctx() -> LoopContext:
    return LoopContext(config=LoopConfig())


def _document(path):
    document = docx.Document()
    title = document.add_paragraph("Template Heading")
    title.style = "Title"
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Section 1")
    document.add_paragraph("Revenue is 10.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    document.save(path)


def test_word_read_reports_paragraphs_tables_and_tracks_guard(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()

    result = asyncio.run(WordReadTool().run({"path": str(target)}, ctx))

    assert result.is_error is False
    data = json.loads(result.content)
    assert data["type"] == "word_read"
    assert data["paragraph_count"] == 3
    assert data["paragraphs"][0]["text"] == "Template Heading"
    assert data["tables"][0]["cells"][0]["text"] == "Metric"
    assert str(target.resolve()) in ctx.scratch["word_read_files"]


def test_word_edit_requires_word_read_first(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)

    result = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {
                        "op": "replace_text",
                        "paragraph_index": 2,
                        "old": "10",
                        "new": "12",
                    }
                ],
            },
            _ctx(),
        )
    )

    assert result.is_error is True
    assert "Call WordRead first" in result.content
    assert "Revenue is 10" in "\n".join(p.text for p in docx.Document(target).paragraphs)


def test_word_edit_applies_scoped_replace_and_rejects_global_by_default(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx))

    global_result = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [{"op": "replace_text", "old": "Revenue", "new": "Sales"}],
            },
            ctx,
        )
    )
    scoped_result = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {
                        "op": "replace_text",
                        "paragraph_index": 2,
                        "old": "10",
                        "new": "12",
                    }
                ],
            },
            ctx,
        )
    )

    assert global_result.is_error is True
    assert "paragraph_index" in global_result.content
    assert scoped_result.is_error is False
    updated = docx.Document(target)
    assert updated.paragraphs[2].text == "Revenue is 12."
    assert str(target.resolve()) in ctx.scratch["edited_files"]


def test_word_edit_copies_paragraph_style_and_inserts_local_paragraph(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx))

    result = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {
                        "op": "copy_paragraph_style",
                        "source_paragraph_index": 0,
                        "paragraph_index": 1,
                    },
                    {
                        "op": "insert_paragraph_after",
                        "paragraph_index": 1,
                        "text": "Inserted local note.",
                        "style": "Normal",
                    },
                ],
            },
            ctx,
        )
    )

    assert result.is_error is False
    updated = docx.Document(target)
    assert updated.paragraphs[1].style.name == updated.paragraphs[0].style.name
    assert updated.paragraphs[1].runs[0].bold is True
    assert updated.paragraphs[2].text == "Inserted local note."
    assert updated.paragraphs[3].text == "Revenue is 10."


def test_word_edit_supports_thesis_structure_ops(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx))

    result = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {"op": "set_heading_level", "paragraph_index": 1, "level": 1},
                    {
                        "op": "insert_table_after",
                        "paragraph_index": 2,
                        "rows": [["Sample", "Temp"], ["A", "25"]],
                    },
                    {
                        "op": "insert_toc_after",
                        "paragraph_index": 0,
                        "levels": "1-3",
                    },
                    {
                        "op": "add_footnote",
                        "paragraph_index": 2,
                        "text": "A concise explanatory note.",
                    },
                    {
                        "op": "set_header",
                        "text": "Thesis Title",
                    },
                    {
                        "op": "set_footer",
                        "text": "Page",
                        "page_number": True,
                    },
                ],
            },
            ctx,
        )
    )

    assert result.is_error is False
    updated = docx.Document(target)
    section = next(p for p in updated.paragraphs if p.text == "Section 1")
    assert section.style.name == "Heading 1"
    assert len(updated.tables) == 2
    assert updated.tables[0].cell(0, 0).text == "Sample"
    assert "Thesis Title" in updated.sections[0].header.paragraphs[0].text
    assert "Page" in updated.sections[0].footer.paragraphs[0].text
    with zipfile.ZipFile(target) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        footer_xml = zf.read("word/footer1.xml").decode("utf-8")
        footnotes_xml = zf.read("word/footnotes.xml").decode("utf-8")
    assert "TOC" in document_xml
    assert "PAGE" in footer_xml
    assert "A concise explanatory note." in footnotes_xml


def test_word_read_reports_document_structure(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx))
    asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {"op": "set_heading_level", "paragraph_index": 1, "level": 2},
                    {"op": "set_header", "text": "Header"},
                    {"op": "set_footer", "text": "Footer", "page_number": True},
                    {"op": "add_footnote", "paragraph_index": 2, "text": "Note"},
                    {"op": "insert_toc_after", "paragraph_index": 0},
                ],
            },
            ctx,
        )
    )

    read_result = asyncio.run(WordReadTool().run({"path": str(target)}, ctx))

    data = json.loads(read_result.content)
    structure = data["structure"]
    assert structure["headings"][0]["text"] == "Section 1"
    assert structure["headings"][0]["style"] == "Heading 2"
    assert structure["has_toc_field"] is True
    assert structure["has_footnotes_part"] is True
    assert structure["footnotes"][0]["text"] == "Note"
    assert structure["headers"][0]["text"] == "Header"
    assert "Footer" in structure["footers"][0]["text"]
    assert structure["has_page_field"] is True


def test_word_edit_rejects_accidental_empty_text(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx))

    result = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [{"op": "set_paragraph_text", "paragraph_index": 1, "text": ""}],
            },
            ctx,
        )
    )

    assert result.is_error is True
    assert "would clear existing text" in result.content
    assert docx.Document(target).paragraphs[1].text == "Section 1"


def test_word_edit_allows_explicit_op_level_empty_text(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx))

    result = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {
                        "op": "set_paragraph_text",
                        "paragraph_index": 1,
                        "text": "",
                        "allow_empty_text": True,
                    }
                ],
            },
            ctx,
        )
    )

    assert result.is_error is False
    assert docx.Document(target).paragraphs[1].text == ""


def test_word_edit_can_add_footnote_after_reload(tmp_path):
    target = tmp_path / "sample.docx"
    _document(target)
    ctx = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx))
    first = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {
                        "op": "add_footnote",
                        "paragraph_index": 2,
                        "text": "First note.",
                    }
                ],
            },
            ctx,
        )
    )
    assert first.is_error is False

    ctx2 = _ctx()
    asyncio.run(WordReadTool().run({"path": str(target)}, ctx2))
    second = asyncio.run(
        WordEditTool().run(
            {
                "path": str(target),
                "ops": [
                    {
                        "op": "add_footnote",
                        "paragraph_index": 2,
                        "text": "Second note.",
                    }
                ],
            },
            ctx2,
        )
    )

    assert second.is_error is False
    with zipfile.ZipFile(target) as zf:
        footnotes_xml = zf.read("word/footnotes.xml").decode("utf-8")
    assert "First note." in footnotes_xml
    assert "Second note." in footnotes_xml


def test_word_tool_protocol_flags_are_minimal():
    tools = word_toolset()

    assert set(tools) == {"WordRead", "WordEdit"}
    assert tools["WordRead"].permission_level == PermissionLevel.SAFE
    assert tools["WordEdit"].permission_level == PermissionLevel.NEEDS_APPROVAL
    assert tools["WordRead"].parallel_safe is True
    assert tools["WordEdit"].parallel_safe is False
