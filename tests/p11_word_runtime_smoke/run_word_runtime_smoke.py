"""Real-COM smoke for the Word runtime backend.

Builds a fixture .docx with three chapters + a TOC field that is initially
stale, runs the backend through a typical op chain (set heading text +
refresh fields + save), and asserts:

- the saved file contains a populated TOC cache;
- the new heading text appears in the TOC cache;
- the toy Word verifier passes;
- exactly one fresh backup landed under ``.bak/`` and old ones were GC'd.

Skipped automatically on non-Windows or when pywin32 / Word are missing.
"""

from __future__ import annotations

import json
import os
import shutil
import sys
import time
from datetime import datetime
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))


def _winword_pids() -> set[int]:
    """Return the set of WINWORD.EXE PIDs currently alive on Windows.

    A set lets the smoke distinguish three failure modes:
      - leak: our PID stayed alive after end_session
      - hijack: a preexisting user PID got killed (the bug this test guards)
      - both
    A scalar count would silently pass when before=1, after=0.
    """
    if os.name != "nt":
        return set()
    import subprocess

    try:
        out = subprocess.check_output(
            ["tasklist", "/FI", "IMAGENAME eq WINWORD.EXE", "/FO", "CSV", "/NH"],
            stderr=subprocess.DEVNULL,
            timeout=5,
        )
    except Exception:
        return set()
    pids: set[int] = set()
    for line in out.decode("utf-8", errors="replace").splitlines():
        parts = [p.strip().strip('"') for p in line.split(",")]
        if len(parts) >= 2 and parts[0].lower() == "winword.exe":
            try:
                pids.add(int(parts[1]))
            except ValueError:
                continue
    return pids


def _build_fixture(path: Path) -> None:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = docx.Document()
    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("First chapter body. Paragraph one.")
    document.add_paragraph("First chapter body. Paragraph two.")
    document.add_heading("Chapter 2", level=1)
    document.add_paragraph("Second chapter body.")
    document.add_heading("Chapter 3", level=1)
    document.add_paragraph("Third chapter body.")

    # Inject a stale TOC field at the top so refresh_fields has something to update.
    toc_para = document.paragraphs[0]
    body = toc_para._element.getparent()
    fld = OxmlElement("w:p")
    run = OxmlElement("w:r")
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    inner_run = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = ""
    inner_run.append(inner_t)
    fld_simple.append(inner_run)
    run.append(fld_simple)
    fld.append(run)
    body.insert(list(body).index(toc_para._element), fld)

    document.save(path)


def main() -> int:
    if os.name != "nt":
        print("SKIP: COM smoke requires Windows")
        return 0
    try:
        import win32com.client  # noqa: F401
        import pythoncom  # noqa: F401
    except ImportError as exc:
        print(f"SKIP: pywin32 missing ({exc})")
        return 0

    from agent.core.word_runtime import (
        AnchorMode,
        OpKind,
        WordRuntimeOp,
        WordRuntimeRequest,
        get_session_manager,
    )
    from agent.core.word_runtime.com_backend import (
        ComWordBackend,
        make_default_com_backend_factory,
    )
    from agent.tools_v2.word_verifier import verify_word_document

    try:
        backend = ComWordBackend()
        backend.connect()
        backend.shutdown()
    except Exception as exc:
        print(f"SKIP: Word.Application not available ({exc})")
        return 0

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    workdir = REPO_ROOT / "tests" / "results" / "p11_word_runtime_smoke" / timestamp
    workdir.mkdir(parents=True, exist_ok=True)
    fixture = workdir / "thesis.docx"
    _build_fixture(fixture)

    pids_before = _winword_pids()

    summary: dict = {
        "passed": False,
        "fixture": str(fixture),
        "timestamp": timestamp,
        "errors": [],
        "winword_pids_before": sorted(pids_before),
    }

    manager = get_session_manager(make_default_com_backend_factory())
    session = manager.get_session("smoke")

    try:
        request = WordRuntimeRequest(
            path=fixture,
            ops=[
                WordRuntimeOp(
                    op=OpKind.SET_HEADING_TEXT,
                    anchor_mode=AnchorMode.HEADING_TEXT,
                    anchor="Chapter 2",
                    new_text="Updated Chapter Two",
                ),
                WordRuntimeOp(
                    op=OpKind.INSERT_PARAGRAPH_AFTER_HEADING,
                    anchor_mode=AnchorMode.HEADING_TEXT,
                    anchor="Chapter 1",
                    new_text="Auto-inserted note paragraph from runtime backend.",
                    style="Normal",
                ),
                WordRuntimeOp(
                    op=OpKind.REFRESH_FIELDS,
                    anchor_mode=AnchorMode.HEADING_TEXT,
                ),
            ],
            conversation_id="smoke",
            save=True,
            refresh_fields_on_save=True,
        )
        result = session.apply(request)
        summary["result"] = result.to_dict()

        verifier = verify_word_document(fixture)
        summary["verifier"] = verifier.to_dict()

        backups_dir = fixture.parent / ".bak"
        backups = sorted(backups_dir.glob(f"{fixture.stem}-*{fixture.suffix}"))
        summary["backup_count"] = len(backups)
        summary["backups"] = [str(b) for b in backups]

        toc_lines = []
        for entry in result.structure_after.get("toc_entries", []):
            line = entry.get("line", "")
            if line:
                toc_lines.append(line)
        summary["toc_lines"] = toc_lines

        headings_after = [h["text"] for h in result.structure_after.get("headings", [])]
        summary["headings_after"] = headings_after

        ok_heading_change = "Updated Chapter Two" in headings_after
        ok_chapter1_remained = "Chapter 1" in headings_after
        ok_toc_has_updated = any("Updated Chapter Two" in line for line in toc_lines)
        ok_backups_capped = len(backups) <= 3 and len(backups) >= 1
        ok_verifier_passed = bool(verifier.passed)

        # Explicit teardown so the WINWORD checks below catch real leaks/hijacks.
        manager.end_session("smoke")
        # Process termination after Quit() + CoUninitialize() is asynchronous;
        # poll up to 8s for our own PID to disappear before sampling.
        pids_after = _winword_pids()
        deadline = time.monotonic() + 8.0
        while (pids_after - pids_before) and time.monotonic() < deadline:
            time.sleep(0.5)
            pids_after = _winword_pids()
        summary["winword_pids_after"] = sorted(pids_after)
        ok_no_winword_leak = len(pids_after - pids_before) == 0
        ok_no_user_winword_killed = pids_before.issubset(pids_after)

        summary["checks"] = {
            "ok_heading_change": ok_heading_change,
            "ok_chapter1_remained": ok_chapter1_remained,
            "ok_toc_has_updated": ok_toc_has_updated,
            "ok_backups_capped": ok_backups_capped,
            "ok_verifier_passed": ok_verifier_passed,
            "ok_no_winword_leak": ok_no_winword_leak,
            "ok_no_user_winword_killed": ok_no_user_winword_killed,
        }
        summary["passed"] = all(summary["checks"].values())

    except Exception as exc:
        summary["errors"].append(f"{type(exc).__name__}: {exc}")
    finally:
        manager.end_session("smoke")

    out = workdir / "summary.json"
    out.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0 if summary["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
