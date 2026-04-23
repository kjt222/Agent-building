from __future__ import annotations

import argparse
import json
import shutil
import time
from pathlib import Path
from typing import Any

import requests
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
ARTIFACT_ROOT = (
    ROOT
    / "tests"
    / "p4_word_complex_validation"
    / "2026-04-23-thesis-short-natural"
)


MODELS = {
    "doubao-code": {
        "profile": "doubao-code",
        "model_override": None,
    },
    "gpt-5.4-mini": {
        "profile": "gpt-5.4",
        "model_override": "gpt-5.4-mini",
    },
}


SCENARIOS = {
    "thesis_all_in_one": {
        "title": "\u6bd5\u4e1a\u8bba\u6587\u4e00\u6b21\u6027\u6574\u7406",
        "prompt": (
            "\u8001\u5e08\u8bf4\u8fd9\u4e2a\u6bd5\u4e1a\u8bba\u6587\u683c\u5f0f\u4e0d"
            "\u884c\uff0c\u5e2e\u6211\u6574\u7406\u4e00\u4e0b\uff1a{path}\n"
            "\u76ee\u5f55\u3001\u4e00\u4e8c\u4e09\u7ea7\u6807\u9898\u3001\u8868\u683c"
            "\u3001\u811a\u6ce8\u3001\u9875\u7709\u9875\u811a\u90fd\u8981\u6709\uff0c"
            "\u6b63\u6587\u610f\u601d\u4e0d\u8981\u6539\u3002"
        ),
    },
    "thesis_review_fix": {
        "title": "\u6bd5\u4e1a\u8bba\u6587\u8bc4\u9605\u6574\u6539",
        "prompt": (
            "\u8fd9\u4efd\u8bba\u6587\u88ab\u8bf4\u5f88\u4e0d\u89c4\u8303\uff0c"
            "\u4f60\u5e2e\u6211\u6539\u6210\u50cf\u6bd5\u4e1a\u8bba\u6587\u7684"
            "\u6837\u5b50\uff1a{path}\n"
            "\u6807\u9898\u5c42\u7ea7\u3001\u76ee\u5f55\u3001\u5bf9\u6bd4\u8868"
            "\u3001\u811a\u6ce8\u3001\u9875\u7709\u9875\u811a\u8fd9\u4e9b\u90fd"
            "\u5e2e\u6211\u5904\u7406\u4e00\u4e0b\u3002"
        ),
    },
}


def make_source_doc(path: Path, scenario_id: str) -> None:
    doc = Document()
    doc.add_paragraph("\u4e8c\u7ef4\u6750\u6599\u970d\u5c14\u6548\u5e94\u6d4b\u8bd5\u5e73\u53f0\u8bbe\u8ba1", style="Title")
    doc.add_paragraph("\u6458\u8981")
    doc.add_paragraph("\u672c\u6587\u8ba8\u8bba\u6837\u54c1\u5236\u5907\u3001\u6d4b\u8bd5\u6d41\u7a0b\u4e0e\u7ed3\u679c\u5206\u6790\u3002")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This draft studies a measurement workflow for two-dimensional materials.")
    doc.add_paragraph("\u76ee\u5f55")
    doc.add_paragraph("\u8fd9\u91cc\u8fd8\u662f\u624b\u5de5\u76ee\u5f55\u5360\u4f4d\u3002")
    doc.add_paragraph("\u7b2c\u4e00\u7ae0 \u7eea\u8bba")
    doc.add_paragraph("1.1 \u7814\u7a76\u80cc\u666f")
    doc.add_paragraph("\u970d\u5c14\u6548\u5e94\u662f\u5224\u65ad\u8f7d\u6d41\u5b50\u7c7b\u578b\u548c\u8fc1\u79fb\u7387\u7684\u91cd\u8981\u65b9\u6cd5\u3002")
    doc.add_paragraph("1.1.1 \u56fd\u5185\u5916\u7814\u7a76\u73b0\u72b6")
    doc.add_paragraph("\u76f8\u5173\u7814\u7a76\u4e3b\u8981\u96c6\u4e2d\u5728\u4f4e\u7ef4\u6750\u6599\u548c\u5668\u4ef6\u7a33\u5b9a\u6027\u3002")
    doc.add_paragraph("\u7b2c\u4e8c\u7ae0 \u5b9e\u9a8c\u65b9\u6cd5")
    doc.add_paragraph("2.1 \u5b9e\u9a8c\u65b9\u6848")
    doc.add_paragraph("\u5b9e\u9a8c\u5305\u542b\u6837\u54c1\u5236\u5907\u3001\u9000\u706b\u5904\u7406\u548c\u7535\u5b66\u6d4b\u8bd5\u3002")
    if scenario_id == "thesis_review_fix":
        doc.add_paragraph("\u8fc1\u79fb\u7387\u53ef\u4ee5\u53cd\u6620\u6750\u6599\u5728\u7535\u573a\u4f5c\u7528\u4e0b\u7684\u8f93\u8fd0\u80fd\u529b\u3002")
    doc.add_paragraph("\u7b2c\u4e09\u7ae0 \u7ed3\u679c\u8ba8\u8bba")
    doc.add_paragraph("3.1 \u6d4b\u8bd5\u7ed3\u679c")
    doc.add_paragraph("\u6d4b\u8bd5\u7ed3\u679c\u663e\u793a\u5668\u4ef6\u5728\u5ba4\u6e29\u4e0b\u5177\u6709\u7a33\u5b9a\u54cd\u5e94\u3002")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "\u5df2\u6709\u9879\u76ee"
    table.cell(0, 1).text = "\u72b6\u6001"
    table.cell(1, 0).text = "\u539f\u59cb\u8868\u683c"
    table.cell(1, 1).text = "\u4fdd\u7559"
    doc.add_paragraph("\u7b2c\u56db\u7ae0 \u7ed3\u8bba")
    doc.add_paragraph("\u672c\u6587\u5b8c\u6210\u4e86\u6d4b\u8bd5\u5e73\u53f0\u7684\u521d\u6b65\u9a8c\u8bc1\u3002")
    doc.add_paragraph("\u53c2\u8003\u6587\u732e")
    doc.add_paragraph("[1] Sample reference.")
    doc.save(path)


def post_select_profile(base_url: str, profile: str) -> None:
    res = requests.post(
        f"{base_url}/profiles/select",
        data={"profile": profile},
        allow_redirects=False,
        timeout=20,
    )
    if res.status_code not in (200, 303):
        raise RuntimeError(f"select profile failed: {res.status_code} {res.text[:200]}")


def create_conversation(base_url: str, profile: str) -> str:
    res = requests.post(
        f"{base_url}/api/conversations",
        json={"profile": profile},
        timeout=20,
    )
    res.raise_for_status()
    return str(res.json()["conversation_id"])


def add_message(
    base_url: str,
    conv_id: str,
    role: str,
    content: str,
    model: str | None = None,
) -> None:
    payload: dict[str, Any] = {"role": role, "content": content}
    if model:
        payload["model"] = model
    res = requests.post(
        f"{base_url}/api/conversations/{conv_id}/messages",
        json=payload,
        timeout=20,
    )
    res.raise_for_status()


def parse_sse(raw_text: str) -> tuple[list[dict], str, dict]:
    events: list[dict] = []
    event_name = "message"
    data_lines: list[str] = []
    assistant_chunks: list[str] = []
    done_payload: dict = {}

    def flush() -> None:
        nonlocal event_name, data_lines, done_payload
        if not data_lines:
            return
        payload = "\n".join(data_lines)
        try:
            data = json.loads(payload)
        except json.JSONDecodeError:
            data = {"raw": payload}
        events.append({"event": event_name, "data": data})
        if event_name == "token":
            assistant_chunks.append(str(data.get("text") or ""))
        if event_name == "done" and isinstance(data, dict):
            done_payload = data
        event_name = "message"
        data_lines = []

    for line in raw_text.splitlines():
        if not line:
            flush()
        elif line.startswith("event:"):
            event_name = line.split(":", 1)[1].strip()
        elif line.startswith("data:"):
            data_lines.append(line.split(":", 1)[1].lstrip())
    flush()
    return events, "".join(assistant_chunks), done_payload


def run_agent(
    base_url: str,
    model_info: dict,
    message: str,
    conv_id: str,
    run_dir: Path,
    max_iterations: int,
    timeout_seconds: float,
) -> tuple[list[dict], str, dict, bool]:
    payload: dict[str, Any] = {
        "message": message,
        "conversation_id": conv_id,
        "history": [],
        "mode": "auto",
        "max_iterations": max_iterations,
    }
    if model_info.get("model_override"):
        payload["model"] = model_info["model_override"]

    raw_path = run_dir / "raw_sse.txt"
    timed_out = False
    started = time.time()
    with requests.post(
        f"{base_url}/api/agent_chat_v2",
        json=payload,
        stream=True,
        timeout=(20, 90),
    ) as res:
        if res.status_code != 200:
            text = res.text
            raw_path.write_text(text, encoding="utf-8")
            return [], "", {"error": text, "status_code": res.status_code}, False
        with raw_path.open("w", encoding="utf-8", newline="\n") as fh:
            for chunk in res.iter_content(chunk_size=4096, decode_unicode=True):
                if chunk:
                    fh.write(chunk)
                    fh.flush()
                if time.time() - started > timeout_seconds:
                    timed_out = True
                    break
    raw_text = raw_path.read_text(encoding="utf-8", errors="replace")
    events, assistant, done = parse_sse(raw_text)
    if timed_out and "error" not in done:
        done = dict(done)
        done["timed_out"] = True
        done["timeout_seconds"] = timeout_seconds
    return events, assistant, done, timed_out


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [
        {
            "index": i,
            "text": p.text,
            "style": p.style.name if p.style else None,
        }
        for i, p in enumerate(doc.paragraphs)
    ]
    tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables]
    import zipfile

    xml: dict[str, str] = {}
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        for name in (
            "word/document.xml",
            "word/footnotes.xml",
            "word/header1.xml",
            "word/footer1.xml",
        ):
            if name in names:
                xml[name] = zf.read(name).decode("utf-8", errors="replace")
    document_xml = xml.get("word/document.xml", "")
    footer_xml = xml.get("word/footer1.xml", "")
    return {
        "paragraphs": paragraphs,
        "table_count": len(doc.tables),
        "tables": tables,
        "header_text": "\n".join(
            p.text for section in doc.sections for p in section.header.paragraphs if p.text
        ),
        "footer_text": "\n".join(
            p.text for section in doc.sections for p in section.footer.paragraphs if p.text
        ),
        "has_toc_field": "TOC" in document_xml or "w:fldSimple" in document_xml,
        "has_footnotes_part": bool(xml.get("word/footnotes.xml")),
        "has_page_field": "PAGE" in document_xml or "PAGE" in footer_xml,
    }


def event_tool_manifest(events: list[dict]) -> list[str]:
    for item in events:
        data = item.get("data") or {}
        if item.get("event") == "activity" and data.get("type") == "tool_manifest":
            meta = data.get("meta") or {}
            tools = meta.get("tools")
            if isinstance(tools, list):
                return [str(t) for t in tools]
    return []


def event_tool_path(events: list[dict]) -> list[str]:
    calls: list[str] = []
    for item in events:
        data = item.get("data") or {}
        if item.get("event") == "activity" and data.get("type") == "tool_call":
            meta = data.get("meta") or {}
            calls.append(str(meta.get("name") or ""))
    return calls


def summarize_checks(inspection: dict) -> dict:
    text = "\n".join(p["text"] for p in inspection["paragraphs"])
    table_text = json.dumps(inspection["tables"], ensure_ascii=False)
    heading_styles = {
        p["text"]: p["style"]
        for p in inspection["paragraphs"]
        if p["text"].startswith(("\u7b2c", "1.", "2.", "3."))
    }
    return {
        "heading_styles": heading_styles,
        "toc_field": inspection["has_toc_field"],
        "table_count": inspection["table_count"],
        "new_table_likely_present": (
            "\u6837\u54c1" in table_text
            or "\u65b9\u6cd5" in table_text
            or "\u5bf9\u6bd4" in table_text
        ),
        "footnotes_part": inspection["has_footnotes_part"],
        "header_text": inspection["header_text"],
        "footer_text": inspection["footer_text"],
        "page_field": inspection["has_page_field"],
        "abstract_preserved": "\u6458\u8981" in text and "Abstract" in text,
        "references_preserved": "\u53c2\u8003\u6587\u732e" in text,
    }


def run_one(
    base_url: str,
    model_label: str,
    scenario_id: str,
    max_iterations: int,
    timeout_seconds: float,
) -> dict:
    model_info = MODELS[model_label]
    scenario = SCENARIOS[scenario_id]
    source_dir = ARTIFACT_ROOT / scenario_id / "source"
    source_dir.mkdir(parents=True, exist_ok=True)
    source_doc = source_dir / f"{scenario_id}.docx"
    make_source_doc(source_doc, scenario_id)

    run_dir = ARTIFACT_ROOT / scenario_id / model_label
    if run_dir.exists():
        shutil.rmtree(run_dir)
    run_dir.mkdir(parents=True, exist_ok=True)
    target_doc = run_dir / f"{scenario_id}_{model_label}.docx"
    shutil.copy2(source_doc, target_doc)
    prompt = scenario["prompt"].format(path=str(target_doc))

    post_select_profile(base_url, model_info["profile"])
    conv_id = create_conversation(base_url, model_info["profile"])
    add_message(base_url, conv_id, "user", prompt)
    started = time.time()
    events, assistant, done, timed_out = run_agent(
        base_url,
        model_info,
        prompt,
        conv_id,
        run_dir,
        max_iterations,
        timeout_seconds,
    )
    elapsed = time.time() - started
    if assistant:
        add_message(
            base_url,
            conv_id,
            "assistant",
            assistant,
            model_info.get("model_override") or model_label,
        )

    (run_dir / "prompt.txt").write_text(prompt, encoding="utf-8")
    (run_dir / "events.jsonl").write_text(
        "\n".join(json.dumps(e, ensure_ascii=False) for e in events),
        encoding="utf-8",
    )
    (run_dir / "assistant.txt").write_text(assistant, encoding="utf-8")

    inspection = inspect_docx(target_doc)
    result = {
        "scenario": scenario_id,
        "scenario_title": scenario["title"],
        "profile": model_info["profile"],
        "model_label": model_label,
        "model_override": model_info.get("model_override"),
        "conversation_id": conv_id,
        "elapsed_seconds": round(elapsed, 2),
        "timed_out": timed_out,
        "done": done,
        "tool_manifest": event_tool_manifest(events),
        "tool_path": event_tool_path(events),
        "inspection": inspection,
        "checks": summarize_checks(inspection),
    }
    (run_dir / "result.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return result


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--base-url", default="http://127.0.0.1:8766")
    parser.add_argument("--model", choices=[*MODELS.keys(), "all"], default="all")
    parser.add_argument("--scenario", choices=[*SCENARIOS.keys(), "all"], default="thesis_all_in_one")
    parser.add_argument("--max-iterations", type=int, default=6)
    parser.add_argument("--timeout-seconds", type=float, default=240)
    args = parser.parse_args()

    base_url = args.base_url.rstrip("/")
    model_labels = list(MODELS) if args.model == "all" else [args.model]
    scenario_ids = list(SCENARIOS) if args.scenario == "all" else [args.scenario]
    ARTIFACT_ROOT.mkdir(parents=True, exist_ok=True)

    results = []
    try:
        for scenario_id in scenario_ids:
            for model_label in model_labels:
                results.append(
                    run_one(
                        base_url,
                        model_label,
                        scenario_id,
                        args.max_iterations,
                        args.timeout_seconds,
                    )
                )
    finally:
        post_select_profile(base_url, "doubao-code")

    summary_path = ARTIFACT_ROOT / "summary.json"
    existing: list[Any] = []
    if summary_path.exists():
        try:
            existing = json.loads(summary_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            existing = []
    summary_path.write_text(
        json.dumps(existing + results, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(results, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
