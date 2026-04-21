from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional


@dataclass
class DocxEditResult:
    replacements: int = 0
    appended: int = 0
    headings: int = 0
    cross_run_merges: int = 0


def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_paragraphs(document) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _replace_within_run(run, old: str, new: str, remaining: Optional[int]) -> int:
    if old not in run.text:
        return 0
    occurrences = run.text.count(old)
    max_count = -1 if remaining is None else min(occurrences, remaining)
    if max_count == 0:
        return 0
    run.text = run.text.replace(old, new, max_count if max_count != -1 else -1)
    return occurrences if max_count == -1 else max_count


def _replace_across_runs(paragraph, old: str, new: str, remaining: Optional[int]) -> tuple[int, int]:
    """Replace `old` across run boundaries while preserving the first spanned
    run's `rPr` (format). Middle/tail runs keep their XML element but lose the
    matched portion of their text; their `rPr` stays intact so later text in
    those runs retains its format.

    Returns (replaced_count, cross_run_merges).
    """
    replaced = 0
    merges = 0
    while remaining is None or remaining > 0:
        runs = list(paragraph.runs)
        if not runs:
            break
        full_text = "".join(r.text for r in runs)
        idx = full_text.find(old)
        if idx == -1:
            break
        end = idx + len(old)

        cursor = 0
        start_run = start_off = None
        end_run = end_off = None
        for i, r in enumerate(runs):
            r_start, r_end = cursor, cursor + len(r.text)
            if start_run is None and r_start <= idx < r_end:
                start_run, start_off = i, idx - r_start
            if r_start < end <= r_end:
                end_run, end_off = i, end - r_start
                break
            cursor = r_end

        if start_run is None or end_run is None:
            break

        if start_run == end_run:
            r = runs[start_run]
            r.text = r.text[:start_off] + new + r.text[end_off:]
        else:
            first = runs[start_run]
            last = runs[end_run]
            # Replacement inherits first run's format; tail of last run keeps
            # its own format in place.
            first.text = first.text[:start_off] + new
            for i in range(start_run + 1, end_run):
                runs[i].text = ""
            last.text = last.text[end_off:]
            merges += 1

        replaced += 1
        if remaining is not None:
            remaining -= 1
    return replaced, merges


def _replace_in_paragraph(
    paragraph,
    old: str,
    new: str,
    remaining: Optional[int],
    fail_on_cross_run: bool,
) -> tuple[int, Optional[int], int]:
    """Replace `old` with `new` in a paragraph, preferring single-run surgical
    edits to preserve formatting. Falls back to run-merge for cross-run matches.
    Returns (replaced, remaining, cross_run_merges).
    """
    if remaining is not None and remaining <= 0:
        return 0, remaining, 0

    replaced = 0
    for run in paragraph.runs:
        if remaining is not None and remaining <= 0:
            break
        count = _replace_within_run(run, old, new, remaining)
        if count:
            replaced += count
            if remaining is not None:
                remaining -= count

    if (remaining is None or remaining > 0) and old in paragraph.text:
        if fail_on_cross_run:
            raise ValueError(
                f"Cross-run match for {old!r} in paragraph; refusing to rewrite "
                "because it would lose run-level formatting. "
                "Use a more specific 'old' string contained in a single run, "
                "or pass fail_on_cross_run=False to merge runs (preserves the "
                "first spanned run's format)."
            )
        across, merges = _replace_across_runs(paragraph, old, new, remaining)
        if across:
            replaced += across
            if remaining is not None:
                remaining -= across
            return replaced, remaining, merges

    return replaced, remaining, 0


def apply_docx_ops(
    path: Path,
    ops: list[dict],
    fail_on_cross_run: bool = False,
) -> DocxEditResult:
    """Apply a sequence of edit ops to a .docx file.

    fail_on_cross_run: when True, replace_text raises if the target text spans
    multiple runs. Default False uses a format-preserving run-merge strategy.
    """
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx not installed. Install `python-docx`.") from exc

    if not path.exists():
        raise FileNotFoundError(path)

    document = docx.Document(str(path))
    result = DocxEditResult()

    for op in ops:
        op_type = op.get("op")
        if op_type == "replace_text":
            old = op.get("old")
            new = op.get("new")
            if old is None or new is None:
                raise ValueError("replace_text requires 'old' and 'new'")
            count = op.get("count", -1)
            remaining = None if count is None or count < 0 else int(count)
            for paragraph in _iter_paragraphs(document):
                replaced, remaining, merges = _replace_in_paragraph(
                    paragraph, str(old), str(new), remaining, fail_on_cross_run
                )
                if replaced:
                    result.replacements += replaced
                if merges:
                    result.cross_run_merges += merges
                if remaining is not None and remaining <= 0:
                    break
        elif op_type == "append_paragraph":
            text = op.get("text", "")
            style = op.get("style")
            document.add_paragraph(str(text), style=style)
            result.appended += 1
        elif op_type == "add_heading":
            text = op.get("text", "")
            level = int(op.get("level", 1))
            document.add_heading(str(text), level=level)
            result.headings += 1
        else:
            raise ValueError(f"Unsupported docx op: {op_type}")

    document.save(str(path))
    return result
