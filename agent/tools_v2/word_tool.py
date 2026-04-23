"""Minimal Word tools for AgentLoop v2.

The Word boundary mirrors the Excel Read/Edit contract. The model can inspect
document structure first, then apply scoped operations. It cannot run arbitrary
python-docx or COM automation scripts through this tool.
"""

from __future__ import annotations

import json
import shutil
import time
import zipfile
from pathlib import Path
from typing import Any

from agent.core.loop import LoopContext, PermissionLevel, ToolResultBlock
from agent.tools.docx_editor import _replace_in_paragraph
from agent.tools_v2.primitives import _ToolBase


_SUPPORTED_SUFFIXES = {".docx"}
_MAX_PARAGRAPHS = 120
_MAX_TABLE_CELLS = 120


def _load_document(path: Path):
    import docx

    return docx.Document(str(path))


def _resolve_word_path(raw_path: Any) -> Path:
    if raw_path is None:
        raise ValueError("path is required")
    path = Path(str(raw_path)).expanduser()
    if not path.exists():
        raise FileNotFoundError(f"file not found: {path}")
    if path.is_dir():
        raise IsADirectoryError(str(path))
    if path.suffix.lower() not in _SUPPORTED_SUFFIXES:
        raise ValueError("expected .docx document")
    return path


def _path_key(path: Path) -> str:
    return str(path.resolve())


def _paragraph_payload(paragraph, index: int, *, include_runs: bool) -> dict:
    payload = {
        "index": index,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
    }
    if include_runs:
        runs = []
        for run_index, run in enumerate(paragraph.runs):
            runs.append({
                "index": run_index,
                "text": run.text,
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "underline": bool(run.underline),
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
                "font_color": getattr(run.font.color, "rgb", None),
            })
        payload["runs"] = runs
    return payload


def _table_payload(table, index: int, *, max_cells: int) -> dict:
    cells = []
    count = 0
    truncated = False
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if count >= max_cells:
                truncated = True
                break
            cells.append({
                "row": row_index,
                "col": col_index,
                "text": cell.text,
            })
            count += 1
        if truncated:
            break
    return {
        "index": index,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "cells": cells,
        "truncated": truncated,
    }


def _document_structure_payload(path: Path) -> dict:
    headings: list[dict[str, Any]] = []
    footnotes: list[dict[str, Any]] = []
    headers: list[dict[str, Any]] = []
    footers: list[dict[str, Any]] = []
    has_toc_field = False
    has_page_field = False
    has_footnotes_part = False

    try:
        document = _load_document(path)
        for idx, paragraph in enumerate(document.paragraphs):
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading"):
                headings.append({
                    "paragraph_index": idx,
                    "text": paragraph.text,
                    "style": style_name,
                })
        for section_index, section in enumerate(document.sections):
            header_text = "\n".join(p.text for p in section.header.paragraphs if p.text)
            footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text)
            headers.append({"section": section_index, "text": header_text})
            footers.append({"section": section_index, "text": footer_text})
    except Exception:
        pass

    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
            document_xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
            has_toc_field = "TOC" in document_xml
            if "word/footer1.xml" in names:
                footer_xml = zf.read("word/footer1.xml").decode("utf-8", errors="replace")
                has_page_field = "PAGE" in footer_xml
            if "word/footnotes.xml" in names:
                has_footnotes_part = True
                from xml.etree import ElementTree as ET

                root = ET.fromstring(zf.read("word/footnotes.xml"))
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for item in root.findall("w:footnote", ns):
                    fid = item.attrib.get(f"{{{ns['w']}}}id")
                    if fid in {"-1", "0"}:
                        continue
                    texts = [node.text or "" for node in item.findall(".//w:t", ns)]
                    footnotes.append({"id": fid, "text": "".join(texts)})
    except Exception:
        pass

    return {
        "headings": headings,
        "headers": headers,
        "footers": footers,
        "has_toc_field": has_toc_field,
        "has_footnotes_part": has_footnotes_part,
        "footnotes": footnotes,
        "has_page_field": has_page_field,
    }


def _copy_paragraph_style(document, source_index: int, target_index: int) -> str:
    source = document.paragraphs[source_index]
    target = document.paragraphs[target_index]
    if source.style:
        target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing

    source_run = next((run for run in source.runs if run.text), None)
    if source_run is not None:
        for run in target.runs:
            run.bold = source_run.bold
            run.italic = source_run.italic
            run.underline = source_run.underline
            run.font.name = source_run.font.name
            run.font.size = source_run.font.size
            source_color = getattr(source_run.font.color, "rgb", None)
            if source_color is not None:
                run.font.color.rgb = source_color
    return f"paragraph_style:{source_index}->{target_index}"


def _set_run_style(paragraph, op: dict) -> int:
    touched = 0
    for run in paragraph.runs:
        if "bold" in op:
            run.bold = bool(op["bold"])
        if "italic" in op:
            run.italic = bool(op["italic"])
        if "underline" in op:
            run.underline = bool(op["underline"])
        if "font_name" in op:
            run.font.name = str(op["font_name"])
        if "font_size_pt" in op:
            from docx.shared import Pt

            run.font.size = Pt(float(op["font_size_pt"]))
        touched += 1
    return touched


def _insert_paragraph_after(paragraph, text: str, style: str | None = None):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    inserted.add_run(text)
    if style:
        inserted.style = style
    return inserted


def _insert_table_after(paragraph, rows: list[list[Any]], style: str | None = None) -> str:
    if not rows:
        raise ValueError("insert_table_after requires non-empty rows")
    column_count = max(len(row) for row in rows)
    if column_count <= 0:
        raise ValueError("insert_table_after requires at least one column")
    document = paragraph.part.document
    table = document.add_table(rows=len(rows), cols=column_count)
    if style:
        table.style = style
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            table.cell(row_index, col_index).text = (
                "" if col_index >= len(row) else str(row[col_index])
            )
    paragraph._p.addnext(table._tbl)
    return f"insert_table_after rows={len(rows)} cols={column_count}"


def _add_field_run(paragraph, instruction: str, placeholder: str | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _insert_toc_after(paragraph, levels: str = "1-3", title: str | None = None) -> str:
    if title:
        title_para = _insert_paragraph_after(paragraph, title, style="Heading 1")
        target = _insert_paragraph_after(title_para, "", style="Normal")
    else:
        target = _insert_paragraph_after(paragraph, "", style="Normal")
    _add_field_run(target, f'TOC \\o "{levels}" \\h \\z \\u', "Update field to generate table of contents.")
    return f"insert_toc_after:{levels}"


def _set_header(document, text: str, section_index: int = 0) -> str:
    section = document.sections[section_index]
    paragraph = section.header.paragraphs[0]
    paragraph.text = text
    return f"set_header:{section_index}"


def _set_footer(document, text: str, *, page_number: bool, section_index: int = 0) -> str:
    section = document.sections[section_index]
    paragraph = section.footer.paragraphs[0]
    paragraph.text = text
    if page_number:
        if paragraph.text:
            paragraph.add_run(" ")
        _add_field_run(paragraph, "PAGE", "1")
    return f"set_footer:{section_index} page_number={page_number}"


def _get_or_create_footnotes_part(document):
    from docx.opc.constants import CONTENT_TYPE as CT
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.opc.part import XmlPart
    from docx.oxml import parse_xml

    try:
        part = document.part.part_related_by(RT.FOOTNOTES)
        if hasattr(part, "element"):
            return part
        xml_part = XmlPart.load(
            part.partname,
            part.content_type,
            part.blob,
            document.part.package,
        )
        for rel in document.part.rels.values():
            if rel.reltype == RT.FOOTNOTES:
                rel._target = xml_part
                break
        return xml_part
    except KeyError:
        xml = (
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            "</w:footnotes>"
        )
        part = XmlPart(
            PackURI("/word/footnotes.xml"),
            CT.WML_FOOTNOTES,
            parse_xml(xml),
            document.part.package,
        )
        document.part.relate_to(part, RT.FOOTNOTES)
        return part


def _add_footnote(paragraph, text: str) -> str:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = _get_or_create_footnotes_part(paragraph.part.document)
    root = part.element
    existing_ids = []
    for child in root:
        raw_id = child.get(qn("w:id"))
        try:
            existing_ids.append(int(raw_id))
        except (TypeError, ValueError):
            continue
    footnote_id = max([0, *existing_ids]) + 1

    footnote = OxmlElement("w:footnote")
    footnote.set(qn("w:id"), str(footnote_id))
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    footnote.append(p)
    root.append(footnote)

    run = paragraph.add_run()
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(footnote_id))
    run._r.append(ref)
    return f"add_footnote:{footnote_id}"


class WordReadTool(_ToolBase):
    name = "WordRead"
    description = (
        "Inspect a .docx document before editing. Returns paragraph indexes, "
        "text, styles, run-level formatting, and table cell text. Use this "
        "before WordEdit."
    )
    input_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "DOCX path"},
            "paragraph_start": {"type": "integer", "default": 0},
            "paragraph_limit": {"type": "integer", "default": 80},
            "include_runs": {"type": "boolean", "default": True},
            "include_tables": {"type": "boolean", "default": True},
        },
        "required": ["path"],
    }
    permission_level = PermissionLevel.SAFE
    parallel_safe = True

    async def run(self, input: dict, ctx: LoopContext) -> ToolResultBlock:
        try:
            path = _resolve_word_path(input.get("path"))
            document = _load_document(path)
            start = max(0, int(input.get("paragraph_start", 0)))
            limit = max(1, min(int(input.get("paragraph_limit", 80)), _MAX_PARAGRAPHS))
            include_runs = bool(input.get("include_runs", True))
            include_tables = bool(input.get("include_tables", True))
            paragraphs = [
                _paragraph_payload(p, idx, include_runs=include_runs)
                for idx, p in enumerate(document.paragraphs)
                if start <= idx < start + limit
            ]
            tables = []
            if include_tables:
                tables = [
                    _table_payload(table, idx, max_cells=_MAX_TABLE_CELLS)
                    for idx, table in enumerate(document.tables)
                ]
            result = {
                "type": "word_read",
                "path": str(path),
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "paragraphs": paragraphs,
                "tables": tables,
                "structure": _document_structure_payload(path),
            }
        except Exception as exc:
            return self._err(f"{type(exc).__name__}: {exc}")

        ctx.scratch.setdefault("word_read_files", set()).add(_path_key(path))
        return self._ok(json.dumps(result, ensure_ascii=False, indent=2, default=str))


class WordEditTool(_ToolBase):
    name = "WordEdit"
    description = (
        "Apply scoped structured edits to a .docx document. The document must "
        "have been inspected with WordRead in this AgentLoop run. Prefer "
        "paragraph_index-scoped operations; global text replacement is rejected "
        "unless allow_global=true."
    )
    input_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "DOCX path"},
            "ops": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "op": {
                            "type": "string",
                            "enum": [
                                "replace_text",
                                "set_paragraph_text",
                                "append_paragraph",
                                "add_heading",
                                "set_heading_level",
                                "insert_paragraph_after",
                                "insert_table_after",
                                "insert_toc_after",
                                "add_footnote",
                                "set_header",
                                "set_footer",
                                "copy_paragraph_style",
                                "set_run_style",
                            ],
                        },
                        "paragraph_index": {"type": "integer"},
                        "source_paragraph_index": {"type": "integer"},
                        "old": {"type": "string"},
                        "new": {"type": "string"},
                        "text": {"type": "string"},
                        "style": {"type": "string"},
                        "level": {"type": "integer"},
                        "levels": {"type": "string"},
                        "title": {"type": "string"},
                        "rows": {
                            "type": "array",
                            "items": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                        },
                        "section_index": {"type": "integer"},
                        "page_number": {"type": "boolean"},
                        "bold": {"type": "boolean"},
                        "italic": {"type": "boolean"},
                        "underline": {"type": "boolean"},
                        "font_name": {"type": "string"},
                        "font_size_pt": {"type": "number"},
                    },
                    "required": ["op"],
                },
            },
            "backup": {"type": "boolean", "default": True},
            "allow_global": {"type": "boolean", "default": False},
            "strict": {"type": "boolean", "default": False},
        },
        "required": ["path", "ops"],
    }
    permission_level = PermissionLevel.NEEDS_APPROVAL
    parallel_safe = False

    async def run(self, input: dict, ctx: LoopContext) -> ToolResultBlock:
        try:
            path = _resolve_word_path(input.get("path"))
            if _path_key(path) not in ctx.scratch.setdefault("word_read_files", set()):
                return self._err(
                    f"document was not inspected in this session: {path}. "
                    "Call WordRead first."
                )
            ops = input.get("ops")
            if not isinstance(ops, list) or not ops:
                return self._err("WordEdit requires a non-empty ops list")
            allow_global = bool(input.get("allow_global", False))
            strict = bool(input.get("strict", False))
            allow_empty_text = bool(input.get("allow_empty_text", False))

            backup_path = None
            if bool(input.get("backup", True)):
                stamp = time.strftime("%Y%m%d%H%M%S")
                backup_path = path.with_name(f"{path.stem}.bak-{stamp}{path.suffix}")
                shutil.copy2(path, backup_path)

            document = _load_document(path)
            touched: list[str] = []
            for op in ops:
                if not isinstance(op, dict):
                    raise ValueError("each op must be an object")
                op_type = op.get("op")
                paragraph_index = op.get("paragraph_index")
                paragraph = None
                if paragraph_index is not None:
                    paragraph_index = int(paragraph_index)
                    if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                        raise ValueError(f"paragraph_index out of range: {paragraph_index}")
                    paragraph = document.paragraphs[paragraph_index]

                if op_type == "replace_text":
                    if op.get("old") is None or op.get("new") is None:
                        raise ValueError("replace_text requires 'old' and 'new'")
                    if paragraph is None and not allow_global:
                        raise ValueError(
                            "replace_text requires paragraph_index unless "
                            "allow_global=true"
                        )
                    targets = [paragraph] if paragraph is not None else document.paragraphs
                    total = 0
                    merges = 0
                    for target in targets:
                        replaced, _, cross = _replace_in_paragraph(
                            target,
                            str(op["old"]),
                            str(op["new"]),
                            None,
                            strict,
                        )
                        total += replaced
                        merges += cross
                    touched.append(
                        f"replace_text:{paragraph_index if paragraph is not None else 'global'} "
                        f"replacements={total} cross_run_merges={merges}"
                    )
                elif op_type == "set_paragraph_text":
                    if paragraph is None:
                        raise ValueError("set_paragraph_text requires paragraph_index")
                    text = str(op.get("text") or "")
                    op_allow_empty_text = bool(op.get("allow_empty_text", allow_empty_text))
                    if not text and paragraph.text and not op_allow_empty_text:
                        raise ValueError(
                            "set_paragraph_text would clear existing text; set "
                            "allow_empty_text=true if that is intentional"
                        )
                    if paragraph.runs:
                        paragraph.runs[0].text = text
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        paragraph.add_run(text)
                    touched.append(f"paragraph:{paragraph_index}")
                elif op_type == "append_paragraph":
                    p = document.add_paragraph(str(op.get("text") or ""), style=op.get("style"))
                    touched.append(f"append_paragraph:{len(document.paragraphs) - 1}:{p.style.name}")
                elif op_type == "add_heading":
                    level = int(op.get("level", 1))
                    document.add_heading(str(op.get("text") or ""), level=level)
                    touched.append(f"add_heading:{level}")
                elif op_type == "set_heading_level":
                    if paragraph is None:
                        raise ValueError("set_heading_level requires paragraph_index")
                    level = max(1, min(int(op.get("level", 1)), 9))
                    paragraph.style = f"Heading {level}"
                    if op.get("text") is not None:
                        text = str(op.get("text") or "")
                        op_allow_empty_text = bool(op.get("allow_empty_text", allow_empty_text))
                        if not text and paragraph.text and not op_allow_empty_text:
                            raise ValueError(
                                "set_heading_level text would clear existing text; "
                                "set allow_empty_text=true if intentional"
                            )
                        if paragraph.runs:
                            paragraph.runs[0].text = text
                            for run in paragraph.runs[1:]:
                                run.text = ""
                        else:
                            paragraph.add_run(text)
                    touched.append(f"set_heading_level:{paragraph_index}:{level}")
                elif op_type == "insert_paragraph_after":
                    if paragraph is None:
                        raise ValueError("insert_paragraph_after requires paragraph_index")
                    _insert_paragraph_after(
                        paragraph,
                        str(op.get("text") or ""),
                        op.get("style"),
                    )
                    touched.append(f"insert_after:{paragraph_index}")
                elif op_type == "insert_table_after":
                    if paragraph is None:
                        raise ValueError("insert_table_after requires paragraph_index")
                    rows = op.get("rows")
                    if not isinstance(rows, list):
                        raise ValueError("insert_table_after requires rows")
                    touched.append(_insert_table_after(paragraph, rows, op.get("style")))
                elif op_type == "insert_toc_after":
                    if paragraph is None:
                        raise ValueError("insert_toc_after requires paragraph_index")
                    touched.append(
                        _insert_toc_after(
                            paragraph,
                            str(op.get("levels") or "1-3"),
                            str(op["title"]) if op.get("title") is not None else None,
                        )
                    )
                elif op_type == "add_footnote":
                    if paragraph is None:
                        raise ValueError("add_footnote requires paragraph_index")
                    text = str(op.get("text") or "")
                    if not text:
                        raise ValueError("add_footnote requires text")
                    touched.append(_add_footnote(paragraph, text))
                elif op_type == "set_header":
                    text = str(op.get("text") or "")
                    if not text:
                        raise ValueError("set_header requires text")
                    section_index = int(op.get("section_index") or 0)
                    touched.append(_set_header(document, text, section_index))
                elif op_type == "set_footer":
                    text = str(op.get("text") or "")
                    section_index = int(op.get("section_index") or 0)
                    touched.append(
                        _set_footer(
                            document,
                            text,
                            page_number=bool(op.get("page_number", False)),
                            section_index=section_index,
                        )
                    )
                elif op_type == "copy_paragraph_style":
                    source_index = op.get("source_paragraph_index")
                    if source_index is None or paragraph_index is None:
                        raise ValueError(
                            "copy_paragraph_style requires source_paragraph_index "
                            "and paragraph_index"
                        )
                    source_index = int(source_index)
                    paragraph_index = int(paragraph_index)
                    if source_index < 0 or source_index >= len(document.paragraphs):
                        raise ValueError(f"source_paragraph_index out of range: {source_index}")
                    touched.append(_copy_paragraph_style(document, source_index, paragraph_index))
                elif op_type == "set_run_style":
                    if paragraph is None:
                        raise ValueError("set_run_style requires paragraph_index")
                    count = _set_run_style(paragraph, op)
                    touched.append(f"set_run_style:{paragraph_index} runs={count}")
                else:
                    raise ValueError(f"unsupported WordEdit op: {op_type}")

            document.save(str(path))
            result = {
                "type": "word_edit",
                "path": str(path),
                "backup_path": str(backup_path) if backup_path else None,
                "ops_applied": len(ops),
                "touched": touched,
            }
            ctx.scratch.setdefault("word_edited_files", set()).add(_path_key(path))
            ctx.scratch.setdefault("edited_files", set()).add(_path_key(path))
            return self._ok(json.dumps(result, ensure_ascii=False, indent=2, default=str))
        except Exception as exc:
            return self._err(f"{type(exc).__name__}: {exc}")


def word_toolset() -> dict:
    tools = [WordReadTool(), WordEditTool()]
    return {tool.name: tool for tool in tools}
