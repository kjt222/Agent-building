"""Playwright UI smoke for the WordRuntimeEdit structured diff (P12.2.1).

A scripted adapter calls WordRuntimeEdit with three ops (rename a heading,
insert a paragraph under another heading, refresh fields). The UI should
render a structured ops card with the heading rename shown as a clear
before/after, the insertion under a named anchor, and a refresh_fields row.

Tool execution itself requires a real .docx; the smoke generates one with
python-docx. We then click Accept and verify the heading actually changed
in the saved document.

Skips silently when win32com is unavailable so this runner stays useful
in mixed test environments.
"""

from __future__ import annotations

import argparse
import json
import os
import socket
import subprocess
import sys
import time
import urllib.request
from contextlib import closing
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

RESULTS = ROOT / "tests" / "results" / "p12_word_diff_preview_ui"


def _timestamp() -> str:
    return time.strftime("%Y%m%d_%H%M%S")


def find_free_port() -> int:
    with closing(socket.socket(socket.AF_INET, socket.SOCK_STREAM)) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def wait_for_server(base_url: str, timeout_s: float = 30.0) -> None:
    deadline = time.time() + timeout_s
    last_exc = ""
    while time.time() < deadline:
        try:
            with urllib.request.urlopen(f"{base_url}/api/agent_runtime", timeout=3) as resp:
                if 200 <= resp.status < 300:
                    return
        except Exception as exc:
            last_exc = str(exc)
        time.sleep(0.5)
    raise RuntimeError(f"server did not start in {timeout_s}s: {last_exc}")


def _build_fixture(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading("第一章 绪论", level=1)
    document.add_paragraph("这是第一章正文。")
    document.add_heading("第二章 原理", level=1)
    document.add_paragraph("这是第二章正文。")
    document.add_heading("第三章 结论", level=1)
    document.add_paragraph("这是第三章正文。")
    document.save(path)


def _heading_text(path: Path, target: str) -> str | None:
    import docx

    doc = docx.Document(path)
    for para in doc.paragraphs:
        if (para.style.name or "").startswith("Heading") and target in para.text:
            return para.text
    return None


def _write_config(run_dir: Path) -> Path:
    cfg = run_dir / "config"
    cfg.mkdir(parents=True, exist_ok=True)
    (cfg / "app.yaml").write_text(
        """
active_profile: smoke
profiles:
  smoke: {}
active_kbs: []
knowledge_bases: []
runtime:
  mode: inline
  access_mode: restricted
  monitor:
    enabled: false
    wake_on_task_complete: false
    heartbeat_seconds: 30
""".strip(),
        encoding="utf-8",
    )
    (cfg / "models.yaml").write_text(
        """
profiles:
  smoke:
    llm:
      active: provider
      providers:
        provider:
          type: openai
          model: gpt-5-mini
          api_key_ref: smoke.llm.provider
""".strip(),
        encoding="utf-8",
    )
    return cfg


def _build_fake_adapter_module(run_dir: Path, target_file: Path) -> Path:
    sc_dir = run_dir / "sitepatch"
    sc_dir.mkdir(parents=True, exist_ok=True)
    (sc_dir / "sitecustomize.py").write_text(
        f"""
from agent.core.loop import TextDelta, ToolUseDelta, TurnEnd


_TARGET = r"{str(target_file)}"


class _WordRuntimeAdapter:
    def __init__(self, model, api_key, base_url=None):
        self.calls = 0

    async def stream(self, messages, tools, system=None, **options):
        self.calls += 1
        if self.calls == 1:
            yield ToolUseDelta(
                id='call-wr-1',
                name='WordRuntimeEdit',
                input_partial={{
                    'path': _TARGET,
                    'ops': [
                        {{
                            'op': 'set_heading_text',
                            'anchor_heading': '第二章 原理',
                            'new_text': '第二章 实验方法',
                        }},
                        {{
                            'op': 'insert_paragraph_after_heading',
                            'anchor_heading': '第二章 实验方法',
                            'new_text': '本章已由智能体更新。',
                            'style': 'Normal',
                        }},
                        {{'op': 'refresh_fields'}},
                    ],
                    'save': True,
                    'refresh_fields': True,
                }},
            )
            yield TurnEnd(stop_reason='tool_use', usage={{'total_tokens': 1}})
            return
        yield TextDelta(text='done.')
        yield TurnEnd(stop_reason='end_turn', usage={{'total_tokens': 1}})


def _install():
    import agent.models.openai_responses_adapter as mod
    mod.OpenAIResponsesAdapter = _WordRuntimeAdapter
    import agent.ui.server as srv
    srv.resolve_api_key = lambda **_: 'smoke-key'


_install()
""".strip(),
        encoding="utf-8",
    )
    return sc_dir


def run(args: argparse.Namespace) -> dict[str, Any]:
    # Skip cleanly without COM stack so the runner still finishes on Linux/CI.
    try:
        import win32com.client  # noqa
    except Exception:
        summary = {
            "passed": True,
            "skipped": True,
            "reason": "win32com not available",
            "checks": {},
        }
        out_dir = RESULTS / _timestamp()
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / "summary.json").write_text(
            json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        return 0

    from playwright.sync_api import sync_playwright

    run_dir = RESULTS / _timestamp()
    run_dir.mkdir(parents=True, exist_ok=True)
    target_file = run_dir / "report.docx"
    _build_fixture(target_file)
    cfg_dir = _write_config(run_dir)
    sc_dir = _build_fake_adapter_module(run_dir, target_file)
    port = find_free_port()
    base_url = f"http://127.0.0.1:{port}"

    env = dict(os.environ)
    env["PYTHONPATH"] = (
        str(sc_dir) + os.pathsep + env.get("PYTHONPATH", "")
    )
    env["AGENT_CONFIG_DIR"] = str(cfg_dir)

    server_proc = subprocess.Popen(
        [
            sys.executable,
            "-m", "uvicorn",
            "--factory",
            "--host", "127.0.0.1",
            "--port", str(port),
            "--app-dir", str(ROOT),
            "agent.ui.server:create_app",
        ],
        env=env,
        cwd=str(ROOT),
        stdout=(run_dir / "server_stdout.txt").open("w", encoding="utf-8", errors="replace"),
        stderr=(run_dir / "server_stderr.txt").open("w", encoding="utf-8", errors="replace"),
        text=True,
    )

    summary: dict[str, Any] = {"passed": False, "errors": [], "checks": {}}
    try:
        wait_for_server(base_url)
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_context(viewport={"width": 1400, "height": 900}).new_page()
            page.goto(base_url, wait_until="domcontentloaded")
            page.wait_for_selector("#chat-send", timeout=10_000)

            page.fill("#chat-input", "把这份 Word 文档的第二章标题改成『实验方法』。")
            page.click("#chat-send")
            page.wait_for_selector(".diff-preview-card", timeout=15_000)
            page.wait_for_selector(".diff-preview-op", timeout=5_000)
            card_text = page.locator(".diff-preview-card").last.inner_text(timeout=2_000)
            ops_count = page.locator(".diff-preview-op").count()
            page.screenshot(path=str(run_dir / "01_word_diff_card.png"), full_page=True)

            page.locator(".diff-accept").last.click()
            page.wait_for_selector(".diff-preview-card.is-accepted", timeout=10_000)
            page.wait_for_function(
                "() => Array.from(document.querySelectorAll('.activity-bar')).some(el => ['done','error'].includes(el.dataset.status))",
                timeout=60_000,
            )
            page.screenshot(path=str(run_dir / "02_after_accept.png"), full_page=True)

            browser.close()

        # Validate against the saved .docx.
        time.sleep(0.5)  # allow Word session to flush
        heading_after = _heading_text(target_file, "实验方法")
        summary["card_text"] = card_text
        summary["ops_card_count"] = ops_count
        summary["heading_after"] = heading_after
        # CSS uppercases the op tag; compare case-insensitively for tag labels.
        normalized = card_text.lower()
        summary["checks"] = {
            "card_has_rename_row": "rename heading" in normalized,
            "card_has_insert_row": "insert under heading" in normalized,
            "card_has_refresh_row": "refresh fields" in normalized
            or "refresh_fields" in normalized,
            "card_shows_before_anchor": "第二章 原理" in card_text,
            "card_shows_after_text": "第二章 实验方法" in card_text,
            "heading_actually_renamed": bool(heading_after and "实验方法" in heading_after),
        }
        summary["passed"] = all(summary["checks"].values())
        if not summary["passed"]:
            summary["errors"].append("checks failed: " + json.dumps(summary["checks"]))
    except Exception as exc:
        summary["errors"].append(f"{type(exc).__name__}: {exc}")
    finally:
        server_proc.terminate()
        try:
            server_proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            server_proc.kill()
        (run_dir / "summary.json").write_text(
            json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    print(json.dumps({
        "passed": summary["passed"],
        "errors": summary["errors"],
        "checks": summary.get("checks"),
    }, ensure_ascii=False, indent=2))
    return 0 if summary["passed"] else 1


def main() -> int:
    parser = argparse.ArgumentParser()
    args = parser.parse_args()
    return run(args)


if __name__ == "__main__":
    raise SystemExit(main())
